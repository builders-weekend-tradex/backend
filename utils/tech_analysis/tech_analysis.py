import os
import matplotlib.pyplot as plt
import pandas as pd
import yfinance as yf
import pandas_ta as ta
import matplotlib.dates as mdates
from datetime import datetime, timedelta

from chart_studio import plotly
import plotly
import plotly.offline as pyoff
import plotly.graph_objs as go
import plotly.tools as tls

from docx import Document

TICKER_SYMBOL = 'DIS'

# Define the folder where the files will be stored
output_folder = "utils/tech_analysis/outputs"

# Create the folder if it doesn't exist
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

def tech_analysis(symbol):
        # Define the stock timeframe
        end_date = datetime.today()
        start_date = end_date - timedelta(days=120)  # 4 months before today
        # Fetch stock data using yfinance
        ### TO BE COMMENTED START
        stock_data = yf.download(symbol, start=start_date, end=end_date)
        pd.to_pickle(stock_data, os.path.join(output_folder, "stock_data.pkl"))
        ### TO BE COMMENTED END

        stock_data = pd.read_pickle(os.path.join(output_folder, "stock_data.pkl"))

        # Format DataFrame
        stock_data = stock_data.droplevel(1, axis=1).reset_index()

        # Calculate technical indicators using pandas-ta
        stock_data.ta.macd(append=True)
        stock_data.ta.rsi(append=True)
        stock_data.ta.bbands(append=True)
        stock_data.ta.obv(append=True)

        # Calculate additional technical indicators
        stock_data.ta.sma(length=20, append=True)
        stock_data.ta.ema(length=50, append=True)
        stock_data.ta.stoch(append=True)
        stock_data.ta.adx(append=True)

        # Calculate other indicators
        stock_data.ta.willr(append=True)
        stock_data.ta.cmf(append=True)
        stock_data.ta.psar(append=True)

        #convert OBV to million
        stock_data['OBV_in_million'] =  stock_data['OBV']/1e7
        stock_data['MACD_histogram_12_26_9'] =  stock_data['MACDh_12_26_9'] # not to confuse chatGTP

        # Summarize technical indicators for the last day
        last_day_summary = stock_data.iloc[-1][['Date', 'Close', 'High', 'Low', 'Open', 'Volume', 'MACD_12_26_9',
              'MACDh_12_26_9', 'MACDs_12_26_9', 'RSI_14', 'BBL_5_2.0', 'BBM_5_2.0',
              'BBU_5_2.0', 'BBB_5_2.0', 'BBP_5_2.0', 'OBV', 'SMA_20', 'EMA_50',
              'STOCHk_14_3_3', 'STOCHd_14_3_3', 'ADX_14', 'DMP_14', 'DMN_14',
              'WILLR_14', 'CMF_20', 'PSARl_0.02_0.2', 'PSARs_0.02_0.2',
              'PSARaf_0.02_0.2', 'PSARr_0.02_0.2', 'OBV_in_million',
              'MACD_histogram_12_26_9']]

        # print("Summary of Technical Indicators for the Last Day:")
        # print(last_day_summary)

        ## Work on the prompt
        sys_prompt = """
        Assume the role as a leading Technical Analysis (TA) expert in the stock market, \
        a modern counterpart to Charles Dow, John Bollinger, and Alan Andrews. \
        Your mastery encompasses both stock fundamentals and intricate technical indicators. \
        You possess the ability to decode complex market dynamics, \
        providing clear insights and recommendations backed by a thorough understanding of interrelated factors. \
        Your expertise extends to practical tools like the pandas_ta module, \
        allowing you to navigate data intricacies with ease. \
        As a TA authority, your role is to decipher market trends, make informed predictions, and offer valuable perspectives.

        given {} TA data as below on the last trading day, what will be the next few days possible stock price movement?

        Summary of Technical Indicators for the Last Day:
        {}""".format(symbol,last_day_summary)

        # print(sys_prompt)

        # To create offline GPT Snippet, ##TBD ##VENALI ##POST DEMO
        # from transformers import DistilBertTokenizer, DistilBertModel, GPT2LMHeadModel, GPT2Tokenizer, DeiTFeatureExtractor, DeiTForImageClassification

        # import torch
        # from transformers import DistilBertTokenizer, DistilBertModel, GPT2LMHeadModel, GPT2Tokenizer, CLIPProcessor, CLIPModel

        # # Text Generation with DistilBERT as Encoder and GPT-2 as Decoder
        # def generate_text(prompt):
        #     # Initialize tokenizer and models
        #     bert_tokenizer = DistilBertTokenizer.from_pretrained("distilbert-base-uncased")
        #     bert_model = DistilBertModel.from_pretrained("distilbert-base-uncased")
        #     gpt2_tokenizer = GPT2Tokenizer.from_pretrained("gpt2")
        #     gpt2_model = GPT2LMHeadModel.from_pretrained("gpt2")

        #     # Tokenize the prompt with DistilBERT
        #     inputs = bert_tokenizer(prompt, return_tensors="pt", max_length=1024, truncation=True) # Adjust the parameter as needed

        #     # Get the hidden states from DistilBERT
        #     with torch.no_grad():
        #         outputs = bert_model(**inputs)

        #     # Use the last hidden state as input to GPT-2
        #     input_ids = inputs["input_ids"]
        #     gpt2_inputs = gpt2_tokenizer.encode(prompt, return_tensors="pt", max_length=1024, truncation=True) # Adjust the parameter as needed
        #     gpt2_outputs = gpt2_model.generate(input_ids=gpt2_inputs, max_length=1024, num_beams=5, no_repeat_ngram_size=2, early_stopping=True) # Adjust the parameter as needed

        #     # Decode the output from GPT-2
        #     generated_text = gpt2_tokenizer.decode(gpt2_outputs[0], skip_special_tokens=True)

        #     return generated_text

        # # Let users enter the prompt
        # prompt = input("Enter your prompt: ")
        # generated_text = generate_text(prompt)
        # print("Generated Text:", generated_text)
        generated_text = """Based on the technical snapshot from February 28, 2025, here’s a detailed view of the likely near-term dynamics for DIS:

        Bullish Underpinnings with Cautionary Overbought Signals
        Price and Moving Averages:
        The close at 113.80 is well above both the 20‐day SMA (111.01) and the 50‐day EMA (110.84), reinforcing an overall bullish sentiment. This alignment suggests that, in the short term, buyers have the upper hand.

        MACD & Momentum:
        The MACD line is positive (0.257) with a small but positive histogram (0.272). This indicates that momentum is still in play, although the magnitude isn’t extreme. The MACD supports continued upward movement—if buyers maintain their interest, the trend could persist.

        Oscillator Cues:

        The RSI is around 59.38, a neutral-to-slightly-bullish reading, implying there’s room to run before reaching classic overbought levels.

        However, the stochastic oscillator (with %K near 78.85 and %D around 59.43) is climbing toward levels typically associated with overbought conditions.

        More strikingly, the Williams %R sits at an extreme of –3.02 (with 0 being the extreme overbought territory), which is a clear warning sign that the stock might be due for a short-term pullback or at least a pause as profit-taking sets in.

        Bollinger Bands:
        The price is flirting with the upper Bollinger Band (113.96), suggesting that the stock is at the higher edge of its recent volatility range. In such scenarios, while an upward push is possible, a reversion toward the mid-band (around 111.73) is also common if the current momentum falters.

        Trend Strength & Volume Indicators:
        The ADX at about 17.52 indicates that while the trend is present, it isn’t particularly strong—meaning the upward move could be vulnerable to a brief consolidation or reversal if sentiment shifts. The positive directional movement (DMP > DMN) and a bullish PSAR reading (107.91, well below the current price) continue to support the upward bias, yet the overbought signals inject a level of caution.

        Potential Scenarios for the Next Few Days
        Continued Modest Rally:
        If buying pressure persists and market participants shrug off the overbought warnings, DIS could continue its upward drift, potentially testing resistance levels around 115 or slightly higher. This would be supported by the MACD and the overall price trend above key moving averages.

        Short-Term Correction/Consolidation:
        The overbought readings (especially the Williams %R and near-maximum stochastic levels) suggest that a corrective pullback is possible. In this scenario, profit-taking might drive the price back toward the mid-Bollinger Band (around 111–112) before a fresh bout of buying resumes.

        Sideways Trading:
        Given the moderate ADX and the mix of bullish and overbought signals, another likely outcome is a period of consolidation. The market could settle into a range where prices oscillate between the support near 111 and resistance near 115 until a clearer directional catalyst emerges.

        In Summary
        While the overall setup is bullish—with price above key moving averages and positive momentum indicators—the extreme readings on oscillators and the proximity to the upper Bollinger Band inject a note of caution. Traders should be alert to:

        A break below the SMA/EMA support levels (around 111–112), which could signal the onset of a short-term reversal.
        Signs of weakening momentum on the MACD and a potential divergence in the stochastic, which would further confirm a consolidation phase or a mild pullback.
        Thus, in the coming days, DIS may either continue its moderate rally—if momentum holds—or enter a brief period of consolidation or correction as market participants take profits. Monitoring volume and any shifts in momentum indicators will be key to adjusting positions accordingly.

        This balanced view should serve as a guide for both trend-following and contrarian strategies in the near term.

        """

        print(generated_text)


        ##############################################################################################################################
        # Plot the technical indicators

        # NOTE - If you want to see the charts etc. directly when they are created, please remove "auto_open=false".

        # Price Trend Chart
        price_rend_chart = plt.figure()
        plt.plot(stock_data.index, stock_data['Close'], label='Close', color='blue')
        plt.plot(stock_data.index, stock_data['EMA_50'], label='EMA 50', color='green')
        plt.plot(stock_data.index, stock_data['SMA_20'], label='SMA_20', color='orange')
        plt.title("Price Trend")
        plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%b%d'))  # Format date as "Jun14"
        plt.xticks(rotation=45, fontsize=8)  # Adjust font size
        plotly_fig = tls.mpl_to_plotly(price_rend_chart, resize=True, strip_style=True, verbose=True)
        plotly_fig.update_layout(showlegend=True)
        plotly.offline.plot(plotly_fig, filename=os.path.join(output_folder, 'Price Trend Chart.html'), auto_open=False)


        # On-Balance Volume Chart
        onbalance_volume_chart = plt.figure()
        plt.plot(stock_data['OBV'], label='On-Balance Volume')
        plt.title('On-Balance Volume (OBV) Indicator')
        plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%b%d'))  # Format date as "Jun14"
        plt.xticks(rotation=45, fontsize=8)  # Adjust font size
        plt.legend()
        plotly_fig = tls.mpl_to_plotly(onbalance_volume_chart, resize=True, strip_style=True, verbose=True)
        plotly_fig.update_layout(showlegend=True)
        plotly.offline.plot(plotly_fig, filename=os.path.join(output_folder, 'On-Balance Volume Chart.html'), auto_open=False)

        # MACD Plot
        macd_plot = plt.figure()
        plt.plot(stock_data['MACD_12_26_9'], label='MACD')
        plt.plot(stock_data['MACDh_12_26_9'], label='MACD Histogram')
        plt.title('MACD Indicator')
        plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%b%d'))  # Format date as "Jun14"
        plt.xticks(rotation=45, fontsize=8)  # Adjust font size
        plt.title("MACD")
        plt.legend()
        plotly_fig = tls.mpl_to_plotly(macd_plot, resize=True, strip_style=True, verbose=True)
        plotly_fig.update_layout(showlegend=True)
        plotly.offline.plot(plotly_fig, filename=os.path.join(output_folder, 'MACD.html'), auto_open=False)

        # RSI Plot
        rsi_plot = plt.figure()
        plt.plot(stock_data['RSI_14'], label='RSI')
        plt.axhline(y=70, color='r', linestyle='--', label='Overbought (70)')
        plt.axhline(y=30, color='g', linestyle='--', label='Oversold (30)')
        plt.legend()
        plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%b%d'))  # Format date as "Jun14"
        plt.xticks(rotation=45, fontsize=8)  # Adjust font size
        plotly_fig = tls.mpl_to_plotly(rsi_plot, resize=True, strip_style=True, verbose=True)
        plotly_fig.update_layout(showlegend=True)
        plotly.offline.plot(plotly_fig, filename=os.path.join(output_folder, 'RSI Plot.html'), auto_open=False)

        # Bollinger Bands Plot
        bollinger_bands_plot = plt.figure()
        plt.plot(stock_data.index, stock_data['BBU_5_2.0'], label='Upper BB')
        plt.plot(stock_data.index, stock_data['BBM_5_2.0'], label='Middle BB')
        plt.plot(stock_data.index, stock_data['BBL_5_2.0'], label='Lower BB')
        plt.plot(stock_data.index, stock_data['Close'], label='Close', color='brown')
        plt.title("Bollinger Bands")
        plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%b%d'))  # Format date as "Jun14"
        plt.xticks(rotation=45, fontsize=8)  # Adjust font size
        plt.legend()
        plotly_fig = tls.mpl_to_plotly(bollinger_bands_plot, resize=True, strip_style=True, verbose=True)
        plotly_fig.update_layout(showlegend=True)
        plotly.offline.plot(plotly_fig, filename=os.path.join(output_folder, 'Bollinger Bands Plot.html'), auto_open=False)

        # Stochastic Oscillator Plot
        stochastic_oscillator_plot = plt.figure()
        plt.plot(stock_data.index, stock_data['STOCHk_14_3_3'], label='Stoch %K')
        plt.plot(stock_data.index, stock_data['STOCHd_14_3_3'], label='Stoch %D')
        plt.title("Stochastic Oscillator")
        plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%b%d'))  # Format date as "Jun14"
        plt.xticks(rotation=45, fontsize=8)  # Adjust font size
        plt.legend()
        plotly_fig = tls.mpl_to_plotly(stochastic_oscillator_plot, resize=True, strip_style=True, verbose=True)
        plotly_fig.update_layout(showlegend=True)
        plotly.offline.plot(plotly_fig, filename=os.path.join(output_folder, 'Stochastic Oscillator Plot.html'), auto_open=False)

        # Williams %R Plot
        williams_r_plot = plt.figure()
        plt.plot(stock_data.index, stock_data['WILLR_14'])
        plt.title("Williams %R")
        plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%b%d'))  # Format date as "Jun14"
        plt.xticks(rotation=45, fontsize=8)  # Adjust font size
        plotly_fig = tls.mpl_to_plotly(williams_r_plot, resize=True, strip_style=True, verbose=True)
        plotly_fig.update_layout(showlegend=True)
        plotly.offline.plot(plotly_fig, filename=os.path.join(output_folder, 'Williams R Plot.html'), auto_open=False)

        # ADX Plot
        adx_plot = plt.figure()
        plt.plot(stock_data.index, stock_data['ADX_14'])
        plt.title("Average Directional Index (ADX)")
        plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%b%d'))  # Format date as "Jun14"
        plt.xticks(rotation=45, fontsize=8)  # Adjust font size
        plotly_fig = tls.mpl_to_plotly(adx_plot, resize=True, strip_style=True, verbose=True)
        plotly_fig.update_layout(showlegend=True)
        plotly.offline.plot(plotly_fig, filename=os.path.join(output_folder, 'ADX Plot.html'), auto_open=False)

        # CMF Plot
        cmf_plot = plt.figure()
        plt.plot(stock_data.index, stock_data['CMF_20'])
        plt.title("Chaikin Money Flow (CMF)")
        plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%b%d'))  # Format date as "Jun14"
        plt.xticks(rotation=45, fontsize=8)  # Adjust font size
        plotly_fig = tls.mpl_to_plotly(cmf_plot, resize=True, strip_style=True, verbose=True)
        plotly_fig.update_layout(showlegend=True)
        plotly.offline.plot(plotly_fig, filename=os.path.join(output_folder, 'CMF.html'), auto_open=False)


        # Show the plots (show is disabled at the moment)
        plt.tight_layout()
        # plt.show()

        return generated_text

generated_text = tech_analysis(TICKER_SYMBOL)

document = Document()
document.add_heading(TICKER_SYMBOL, level=1)
document.add_paragraph(generated_text)
document.save(os.path.join(output_folder, 'Summary.docx'))